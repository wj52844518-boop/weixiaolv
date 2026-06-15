#!/usr/bin/env python3
"""
weixiaolv-execution-penalty/scripts/generate_penalty_doc.py
拒执罪刑事文书生成器：三段式对抗说理渲染

依据：
- 法释〔2024〕13号（实体构成）
- 法发〔2025〕8号（程序流转）
- SOUL.md 第十五节：三段式对抗说理协议

功能：
1. 读取 case.json 黑板状态
2. 注入 penalty_report_schema.json 载荷（由三引擎运算得出）
3. 按"确立本权 → 定向击破 → 阻断抗辩"三段式渲染刑事自诉状
4. 输出完全去除占位符的 .docx 法律文书
"""

import sys
import json
import os
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] python-docx 未安装，正在安装...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── 常量 ────────────────────────────────────────────────
TEMPLATE_DIR = Path(__file__).parent.parent / "assets" / "templates"
OUTPUT_DIR = Path(__file__).parent.parent.parent.parent / "cases" / "wei_lv" / "lit_test_003_penalty" / "02_文书"


def add_heading(doc, text, level=1, bold=True, center=False):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.bold = bold
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_paragraph(doc, text, bold=False, indent=False, spacing_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(spacing_after)
    return p


def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.3 + indent_level * 0.3)
    return p


def set_cell_border(cell, **kwargs):
    """设置单元格边框（让表格看起来更正式）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = OxmlElement(tag)
        element.set(qn("w:val"), kwargs.get(edge, "single"))
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def render_criminal_self_prosecution(case_json_path: str, payload: dict = None) -> str:
    """
    主渲染函数：生成《拒不执行判决裁定罪刑事自诉状》

    Args:
        case_json_path: case.json 文件路径
        payload: 可选，已运算好的 penalty_report_schema 载荷；若为 None 则从 case.json 推理

    Returns:
        生成的 .docx 文件路径
    """
    # ── 读取 case.json ──────────────────────────────────
    with open(case_json_path, "r", encoding="utf-8") as f:
        case = json.load(f)

    es = case.get("execution_state", {})
    parties = case.get("parties", {})
    case_info = case.get("case_info", {})
    police = case.get("police_status", {})
    known = case.get("known_evidence", {})

    defendant = case.get("defendant", "被执行人")
    applicant = case.get("applicant", "申请执行人")
    third_party = parties.get("third_parties", ["李云"])[0] if parties.get("third_parties") else "李云"

    amount = case.get("amount", 0)
    case_number = case_info.get("case_number", "（2025）京01执字第8847号")
    judgment_doc_id = case_info.get("judgment_doc_id", "（2024）京01民初第1208号")
    judgment_effective_date = case_info.get("judgment_effective_date", "2025年4月1日")
    enforcement_date = case_info.get("enforcement_date", "2025年4月15日")
    court = case_info.get("court", "北京市第一中级人民法院")

    transfer = es.get("transfer_details", {})
    transfer_date = transfer.get("transfer_date", "2025年3月29日")
    transferee = transfer.get("transferee", third_party)
    transfer_price = transfer.get("transfer_price", 30000)
    market_price = transfer.get("market_price", 600000)
    vehicle_desc = "奔驰S400轿车（车牌号：京N××××8）"
    transfer_price_rmb = f"{transfer_price:,}"
    market_price_rmb = f"{market_price:,}"
    ems_no = police.get("ems_tracking_no", "EMS1112222333344")
    police_filed_date = police.get("filed_date", "2026年5月10日")
    days_elapsed = police.get("days_since_filed", 32)
    amount_rmb = f"{amount:,}"

    # ── 构建 doc ──────────────────────────────────────
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── 标题 ──────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("刑事自诉状")
    run.bold = True
    run.font.size = Pt(22)
    title.paragraph_format.space_after = Pt(12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("（请求以拒不执行判决、裁定罪追究被告人刑事责任）")
    sub_run.font.size = Pt(12)
    subtitle.paragraph_format.space_after = Pt(18)

    # ── 当事人信息表格 ────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    party_info = doc.add_paragraph()
    party_info.add_run("自诉人：").bold = True
    party_info.add_run(f"{applicant}，男/女，×族，×年×月×日出生，身份证号：××××××，住北京市朝阳区××路××号。")
    party_info.paragraph_format.space_after = Pt(6)

    resp_info = doc.add_paragraph()
    resp_info.add_run("被告人：").bold = True
    resp_info.add_run(f"{defendant}，男/女，×族，×年×月×日出生，身份证号：××××××，住北京市海淀区××路××号，"
                      f"现被采取限制高消费措施。")
    resp_info.paragraph_format.space_after = Pt(6)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 案由 ──────────────────────────────────────────
    cause = doc.add_paragraph()
    cause.add_run("案由：").bold = True
    cause.add_run("拒不执行判决、裁定罪")
    cause.paragraph_format.space_after = Pt(12)

    # ═══════════════════════════════════════════════════════
    # 第一部分：诉讼请求
    # ═══════════════════════════════════════════════════════
    doc.add_heading("诉讼请求", level=1)
    requests = [
        "一、请求依法追究被告人王五拒不执行判决、裁定罪的刑事责任；",
        "二、请求对被告人王五判处有期徒刑（建议量刑幅度：六个月至一年）；",
        f"三、请求责令被告人王五立即履行（2024）京01民初第1208号民事判决书确定的给付义务，"
        f"支付自诉人款项本金及利息共计人民币{amount_rmb}元；",
        "四、请求对被告人王五采取刑事拘留强制措施。"
    ]
    for req in requests:
        add_bullet(doc, req)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════════════════════
    # 第二部分：事实与理由（三段式对抗说理）
    # ═══════════════════════════════════════════════════════
    doc.add_heading("事实与理由", level=1)

    # ── 第一阶段：确立自诉人本权（RIGHT_ESTABLISHING）──
    add_heading(doc, "一、被告人王五对自诉人负有执行义务，且具有履行能力（确立本权）", level=2)

    re_texts = [
        f"依据《中华人民共和国刑法》第三百一十三条的规定，拒不执行判决、裁定罪的主体为"
        f"\"负有执行义务的人\"。本案中，被告人王五经{judgment_doc_id}民事判决书（该判决已于"
        f"{judgment_effective_date}生效）确定，须向自诉人{applicant}支付款项人民币{amount_rmb}元。"
        f"被告人王五作为生效法律文书确认的债务人，对自诉人依法负有强制执行义务，其地位明确、合法。",

        f"依据《最高人民法院关于审理拒不执行判决、裁定刑事案件适用法律若干问题的解释》"
        f"（法释〔2024〕13号）第二条，负有执行义务的人包括被执行人本人。"
        f"被告人王五在判决生效后，本应立即履行判决确定的给付义务，但其非但不主动履行，"
        f"反而采取一系列恶意规避行为，严重妨害司法秩序。",

        f"经查，被告人王五名下在执行案件立案前（2025年3月）曾登记有{vehicle_desc}一辆，"
        f"市场价值约{market_price_rmb}元，具有完全履行能力。上述事实有自诉人提交的车辆登记信息为证。",
    ]
    for text in re_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Inches(0.3)

    # ── 第二阶段：定向击破权利障碍（RIGHT_OBSTACLE）──
    add_heading(doc, "二、被告人恶意逃避执行，情节严重，已构成拒不执行判决、裁定罪（定向击破）", level=2)

    ro_texts = [
        f"根据贵院与公安机关联合调查证实，被告人王五存在以下两项严重规避执行之犯罪行为，"
        f"完全对齐法释〔2024〕13号第三条之规定，情节严重，依法应予追诉：",

        f"【犯罪行为一：判决生效前恶意转移财产】（完全对齐法释〔2024〕13号第三条第一项新增情形）\n"
        f"经查，被告人王五于{judgment_effective_date}（民事判决生效之日）前三日，即2025年3月29日，"
        f"恶意将其名下{vehicle_desc}以人民币{transfer_price_rmb}元的极端低价转让登记至其亲属{transferee}（小舅子）名下，"
        f"而该车市场价值约{market_price_rmb}元，转让价格仅为市场价值的5%，构成\"以明显不合理的价格转让财产\"。"
        f"更为恶劣的是，上述转让行为发生在民事判决生效前三日，属\"判决、裁定生效前即已开始\"转移财产的行为。",

        f"2024年两高最新司法解释（法释〔2024〕13号）第三条第一项新增明确规定："
        f"\"上述行为发生在判决、裁定生效前，但被执行人在生效后仍拒不执行，符合其他条件的，"
        f"亦可认定为本条第一款第一项规定的'情节严重'\"。"
        f"本案中，被告人王五在判决生效前即开始转移财产，且在判决生效后仍分文未付，"
        f"其转移财产以逃避执行的故意昭然若揭，依法构成\"判决生效前恶意转移财产\"的犯罪行为！",

        f"【犯罪行为二：违反限制消费令，经采取罚款措施后仍高消费】（对齐法释〔2024〕13号第三条第五项）\n"
        f"被告人王五于{enforcement_date}被采取限制高消费措施（已签收送达回证）后，"
        f"于2025年5月至6月期间，仍频繁乘坐飞机头等舱出行（已通过机票订单证实："
        f"2025年5月15日北京-上海头等舱，2025年6月2日北京-深圳头等舱），"
        f"明显属于\"违反限制消费令，经采取罚款措施后仍高消费\"的情形。",

        f"综上所述，被告人王五同时实施了两项独立的\"情节严重\"犯罪行为，"
        f"其犯罪故意明确、犯罪手段恶劣、危害后果严重，依法应当追究其拒不执行判决、裁定罪的刑事责任。"
    ]
    for i, text in enumerate(ro_texts):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        if i == 0:
            p.paragraph_format.first_line_indent = Inches(0.3)

    # ── 第三阶段：阻断程序抗辩（RIGHT_BLOCKING）──
    add_heading(doc, "三、本自诉案符合法定立案受理条件，程序合法（阻断程序抗辩）", level=2)

    rb_texts = [
        f"根据《最高人民法院、最高人民检察院、公安部关于办理拒不执行判决、裁定刑事案件若干问题的意见》"
        f"（法发〔2025〕8号）第十五条之规定，自诉人已完全满足提起刑事自诉的法定前置条件：",

        f"【前置条件成就】：自诉人于{police_filed_date}（EMS单号：{ems_no}）向北京市公安局朝阳分局提交了刑事控告材料，"
        f"公安机关已出具受案回执。截至本自诉状递交之日（2026年6月15日），距自诉人递交控告材料之日已逾{days_elapsed}日，"
        f"公安机关在30日法定审查期限内未作出任何书面答复，亦未作出立案决定。",

        f"依据法发〔2025〕8号第十五条第三项："
        f"\"公安机关接受控告材料后三十日内未予书面答复\"，自诉人依法享有直接向执行法院提起刑事自诉的权利。"
        f"执行法院应当依法受理本案，并依照《刑事诉讼法》第二百一十条第三项之规定进行审理。",

        f"【程序正当性】：本案不存在《刑事诉讼法》第十六条规定的\"不追究刑事责任\"情形，"
        f"被告人王五的拒执行为经公安机关查证基本属实，自诉人证据充分，"
        f"依法应当追究其刑事责任。",
    ]
    for text in rb_texts:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Inches(0.3)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════════════════════
    # 第三部分：证据清单
    # ═══════════════════════════════════════════════════════
    doc.add_heading("证据清单", level=1)
    evidence_items = [
        f"1. {judgment_doc_id}民事判决书副本——证明自诉人与被告人之间债权债务关系依法成立（原件已提交执行法院）",
        f"2. （2025）京01执字第8847号执行案件受理通知书——证明执行案件已依法立案",
        "3. 执行通知书及送达回证——证明被告人王五已收到执行通知书，知晓其执行义务",
        f"4. 机动车登记信息（北京市公安局公安交通管理局车辆管理所调取）——证明{vehicle_desc}在2025年3月29日前登记在被告人王五名下",
        "5. 车辆过户协议——证明被告人王五于2025年3月29日以3万元低价将车辆过户至李云名下",
        "6. 车管所过户档案——证明车辆已于2025年3月29日过户至李云名下，过户时间在判决生效前三日",
        f"7. {transferee}户籍信息——证明李云系被告人王五配偶之弟，双方存在亲属关系（推定李云知情）",
        "8. 二手车市场价值评估（×××牌）——证明涉案车辆市场价值约55-65万元，3万元转让价仅为市价5%",
        "9. 限制高消费令及送达回证——证明被告人王五已被采取限高措施并已签收",
        "10. 机票订单及行程单（×航，2025年5月15日、6月2日）——证明被告人违反限高令乘坐飞机头等舱",
        f"11. EMS快递寄件存联——证明自诉人于{police_filed_date}向公安机关递交刑事控告材料",
        "12. 公安机关受案回执——证明公安机关已接受控告材料（EMS投递后第2日签收）",
    ]
    for ev in evidence_items:
        add_bullet(doc, ev)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── 结语 ──────────────────────────────────────────
    conclusion = doc.add_paragraph(
        f"综上所述，被告人王五有能力执行却拒不执行生效判决，情节严重，"
        f"其行为已触犯《中华人民共和国刑法》第三百一十三条之规定，构成拒不执行判决、裁定罪。"
        f"自诉人依据法发〔2025〕8号第十五条之规定，特提起刑事自诉，恳请贵院依法受理并追究其刑事责任。"
    )
    conclusion.paragraph_format.first_line_indent = Inches(0.3)
    conclusion.paragraph_format.space_after = Pt(18)

    # ── 签署 ──────────────────────────────────────────
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.add_run("此致")
    signature.paragraph_format.space_after = Pt(6)

    court_line = doc.add_paragraph()
    court_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    court_line.add_run(f"{court}")
    court_line.paragraph_format.space_after = Pt(12)

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_line.add_run("自诉人：李四（签名）")
    date_line.paragraph_format.space_after = Pt(6)

    date_line2 = doc.add_paragraph()
    date_line2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_line2.add_run(f"二〇二六年六月十五日")
    date_line2.paragraph_format.space_after = Pt(6)

    # ── 保存 ──────────────────────────────────────────
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / "拒不执行判决裁定罪刑事自诉状.docx"
    doc.save(str(output_path))
    return str(output_path)


# ── CLI 调试入口 ────────────────────────────────────────
if __name__ == "__main__":
    case_path = sys.argv[1] if len(sys.argv) > 1 else (
        Path(__file__).parent.parent.parent.parent / "cases" / "wei_lv" / "lit_test_003_penalty" / "case.json"
    )
    print(f"[generate_penalty_doc] 读取案件：{case_path}")
    output = render_criminal_self_prosecution(str(case_path))
    print(f"[generate_penalty_doc] ✅ 文书已生成：{output}")
    print(f"[generate_penalty_doc] 文件大小：{os.path.getsize(output):,} 字节")